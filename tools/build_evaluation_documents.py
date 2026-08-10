#!/usr/bin/env python3
"""Build polished Word copies of the versioned AQP evaluation documents."""

from __future__ import annotations

import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, "/root/.codex/skills/builtins/documents/scripts")
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x59, 0x63, 0x6E)
HEADER_FILL = "F2F4F7"
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


DOCUMENTS = [
    {
        "source": ROOT / "docs/CLAIM-NAMING-DECISION-v1.md",
        "output": OUT / "AQP_Claim_and_Naming_Decision_v1.docx",
        "running": "AQP claim and naming decision",
        "version": "Version 1.0 | 10 August 2026",
        "landscape": False,
        "table_font": 9.25,
        "kind": "protocol",
    },
    {
        "source": ROOT / "docs/AQP-FINAL-CONTRIBUTION-v2.md",
        "output": OUT / "AQP_Final_Contribution_Statement_v2.docx",
        "running": "AQP final contribution statement",
        "version": "Version 2.2 | 10 August 2026",
        "landscape": False,
        "table_font": 9.25,
        "kind": "protocol",
    },
    {
        "source": ROOT / "docs/AQP-EVALUATION-MATRIX-v6.md",
        "output": OUT / "AQP_Evaluation_Matrix_v6.docx",
        "running": "AQP evaluation matrix",
        "version": "Version 6.2 | 10 August 2026",
        "landscape": True,
        "table_font": 8.0,
        "kind": "matrix",
    },
    {
        "source": ROOT / "docs/TECHNICAL-EVALUATION-PROTOCOL-v1.0.md",
        "output": OUT / "AQP_Technical_Evaluation_Protocol_v1.docx",
        "running": "AQP technical evaluation",
        "version": "Protocol v1.0 | 9 August 2026",
        "landscape": False,
        "table_font": 9.25,
        "kind": "protocol",
    },
    {
        "source": ROOT / "docs/OBSERVED-RESEARCHER-STUDY-PROTOCOL-v1.0.md",
        "output": OUT / "AQP_Observed_Researcher_Study_Protocol_v1.docx",
        "running": "AQP observed researcher study",
        "version": "Protocol v1.1 | ethics amendment required",
        "landscape": False,
        "table_font": 9.25,
        "kind": "protocol",
    },
    {
        "source": ROOT / "docs/manual-audit/AQP-MANUAL-AT-AUDIT-v1.0.md",
        "output": OUT / "AQP_Manual_AT_Audit_v1.docx",
        "running": "AQP manual assistive-technology audit",
        "version": "Audit v1.1 | not yet executed",
        "landscape": True,
        "table_font": 7.5,
        "kind": "audit",
    },
    {
        "source": ROOT / "docs/SUPERVISOR-FEEDBACK-CLOSURE-MATRIX-v2.md",
        "output": OUT / "AQP_Supervisor_Feedback_Closure_Matrix_v2.docx",
        "running": "AQP supervisor feedback closure matrix",
        "version": "Version 2.1 | 10 August 2026",
        "landscape": True,
        "table_font": 7.5,
        "kind": "matrix",
    },
]


def set_font(run, name: str = "Calibri", size: float | None = None,
             color: RGBColor | None = None, bold: bool | None = None,
             italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, *, size: float, color: RGBColor, bold: bool,
              before: float, after: float, line: float = 1.0) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def find_style(document: Document, name: str):
    """Pandoc uses lowercase built-in w:name values; python-docx exposes aliases."""
    for style in document.styles:
        if style.name == name:
            return style
    raise KeyError(name)


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def widths_for(table, content_width: int, kind: str) -> list[int]:
    columns = len(table.columns)
    if kind == "matrix" and columns == 6:
        return column_widths_from_weights([0.45, 1.65, 2.2, 2.1, 1.45, 1.95], content_width)
    if kind == "audit" and columns == 7:
        return column_widths_from_weights([0.35, 0.85, 1.45, 2.25, 1.65, 1.65, 1.65], content_width)
    if kind == "audit" and columns == 6:
        return column_widths_from_weights([1.1, 1.1, 1.1, 1.1, 1.1, 2.4], content_width)
    if kind == "audit" and columns == 5:
        return column_widths_from_weights([0.8, 1.6, 2.8, 1.4, 0.8], content_width)
    if kind == "risk" and columns == 5:
        return column_widths_from_weights([1.55, 1.1, 0.75, 3.3, 1.7], content_width)
    if kind == "receipt" and columns == 4:
        return column_widths_from_weights([1.65, 1.35, 1.25, 2.25], content_width)
    if columns == 2:
        return column_widths_from_weights([2.0, 4.5], content_width)
    if columns == 3:
        return column_widths_from_weights([2.0, 2.2, 2.3], content_width)
    if columns == 4:
        return column_widths_from_weights([1.2, 1.8, 1.8, 1.7], content_width)
    return column_widths_from_weights([1.0] * columns, content_width)


def format_document(path: Path, *, running: str, version: str,
                    landscape: bool, table_font: float, kind: str) -> None:
    document = Document(path)
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = find_style(document, "Normal")
    set_style(normal, size=11, color=RGBColor(0, 0, 0), bold=False,
              before=0, after=6, line=1.10)
    set_style(find_style(document, "Heading 1"), size=16, color=BLUE, bold=True,
              before=16, after=8)
    set_style(find_style(document, "Heading 2"), size=13, color=BLUE, bold=True,
              before=12, after=6)
    set_style(find_style(document, "Heading 3"), size=12, color=DARK_BLUE, bold=True,
              before=8, after=4)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run(f"{running}  |  {version}")
    set_font(run, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER if landscape else WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    add_field(run, "PAGE")

    # Use one page style on every page. LibreOffice can otherwise create a
    # distinct left/even style with inherited orientation or furniture, which
    # has produced rotated or clipped even pages in rendered QA.
    document.settings.odd_and_even_pages_header_footer = False
    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    ehp = even_header.paragraphs[0]
    ehp.text = ""
    ehp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ehp.paragraph_format.space_after = Pt(0)
    run = ehp.add_run(f"{running}  |  {version}")
    set_font(run, size=8.5, color=MUTED)

    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    efp = even_footer.paragraphs[0]
    efp.text = ""
    efp.alignment = WD_ALIGN_PARAGRAPH.CENTER if landscape else WD_ALIGN_PARAGRAPH.RIGHT
    run = efp.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    add_field(run, "PAGE")

    # LibreOffice 25 currently misplaces left/even running furniture in
    # landscape DOCX pages. These landscape deliverables already repeat their
    # table headings and section labels, so omit unstable decorative furniture
    # instead of shipping clipped text. Do not request separate even-page styles
    # for landscape documents: LibreOffice can otherwise detach a table's first
    # header row on later even pages. Portrait documents retain them.
    if landscape:
        document.settings.odd_and_even_pages_header_footer = False
        for part in (header, footer, even_header, even_footer):
            part.paragraphs[0].text = ""

    first_content = next((p for p in document.paragraphs if p.text.strip()), None)
    if first_content is not None:
        first_content.paragraph_format.space_before = Pt(8)
        first_content.paragraph_format.space_after = Pt(4)
        for run in first_content.runs:
            set_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        # Deliberately start each dense landscape table block on a fresh page.
        # This prevents table continuations from colliding with running page
        # furniture in both Microsoft Word and LibreOffice renderers.
        audit_page_starts = {
            "Frozen test environments",
            "State-by-state audit record",
            "Item operation and voice availability (A05–A09)",
            "Voice listening and safety (A10–A13)",
            "Recovery, review and submission (A14–A18)",
            "Route-wide and imported-scale checks (A19–A23)",
            "Remaining reachable runner and recovery states (A24–A33)",
            "WCAG result summary",
            "Input, error and status criteria",
            "Quantified audit summary",
        }
        if kind == "audit" and paragraph.text in audit_page_starts:
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_before = Pt(42)
        if kind == "matrix" and paragraph.text == "Active claims (continued)":
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_before = Pt(42)
        if paragraph.text.startswith("Status:") or paragraph.text.startswith("Evaluation type:"):
            for run in paragraph.runs:
                set_font(run, size=10, color=MUTED)
        if running == "AQP participant information sheet" and paragraph.text == "How will my data be protected?":
            paragraph.paragraph_format.page_break_before = True
        if paragraph.style.name == "Block Text":
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            for run in paragraph.runs:
                set_font(run, size=10.25, color=DARK_BLUE, italic=True)
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
        for run in paragraph.runs:
            if run.font.name is None:
                set_font(run)

    content_width = section_content_width_dxa(section)
    for table in document.tables:
        table.style = "Table"
        # Keep column meaning visible when a dense audit or matrix table flows
        # onto a second page. The forced section starts prevent a header from
        # being stranded under unrelated content, while the repeat flag keeps
        # continuation rows independently understandable.
        repeat_header(table.rows[0])
        for row_idx, row in enumerate(table.rows):
            prevent_row_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if row_idx == 0:
                    shade_cell(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_font(run, size=table_font, bold=True if row_idx == 0 else None)
        apply_table_geometry(
            table,
            widths_for(table, content_width, kind),
            table_width_dxa=content_width,
            indent_dxa=CELL_MARGINS["start"],
            cell_margins_dxa=CELL_MARGINS,
        )

    document.core_properties.title = first_content.text if first_content else running
    document.core_properties.subject = version
    document.core_properties.author = "Yurui Wang"
    document.core_properties.keywords = "AQP, evaluation, accessibility, questionnaire"
    document.save(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for spec in DOCUMENTS:
        spec["output"].parent.mkdir(parents=True, exist_ok=True)
        subprocess.run(
            [
                "pandoc",
                str(spec["source"]),
                "--from=gfm",
                "--to=docx",
                "--output",
                str(spec["output"]),
            ],
            check=True,
        )
        format_document(
            spec["output"],
            running=spec["running"],
            version=spec["version"],
            landscape=spec["landscape"],
            table_font=spec["table_font"],
            kind=spec["kind"],
        )
        print(spec["output"])


if __name__ == "__main__":
    main()
